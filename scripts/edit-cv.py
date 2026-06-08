"""Apply the website-matching edits to the converted CV docx.

What changes:
  1. "Associate Director/Director" -> "Director"
  2. "(14 total; 1 patent)" -> "(15 total; 3 patents)"
  3. The Loxo bullet "4 direct reports" -> "8 direct reports"
  4. The experience header drops "Loxo Oncology at Eli Lilly" in favor of "Eli Lilly"
     (the website made this same change to streamline the company name)
  5. Add the ACS ECI Award (Aug 2026) as the first row of the Presentations/Awards
     section.

The original cv.original.pdf and the freshly-converted scripts/cv.docx are
preserved as backups; this edits scripts/cv.docx in place.
"""
from __future__ import annotations
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "scripts" / "cv.docx"


def replace_in_runs(paragraph, old: str, new: str) -> int:
    """Replace `old` with `new` across a paragraph's runs.

    python-docx splits text across runs; a phrase can be split across run
    boundaries. So we work on the joined text first, then if a hit exists
    we collapse all runs into the first one with the replacement applied.
    Formatting of the FIRST run wins. That's an acceptable trade-off for
    sentence-level edits.
    """
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return 0

    new_text = full.replace(old, new)
    # Stuff replaced text into first run, blank the rest.
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    return full.count(old)


def main():
    d = Document(str(DOCX))
    counts: dict[str, int] = {}

    # --- Step 1-4: simple text replacements ---------------------------------
    SIMPLE = [
        ("Associate Director/Director", "Director"),
        ("(14 total; 1 patent)", "(15 total; 3 patents)"),
        ("training for 4 direct", "training for 8 direct"),  # Loxo bullet wording
        ("Loxo Oncology at Eli Lilly, Chemistry, Manufacturing, and Controls",
         "Eli Lilly, Chemistry, Manufacturing, and Controls"),
    ]

    for para in d.paragraphs:
        for old, new in SIMPLE:
            n = replace_in_runs(para, old, new)
            counts[old] = counts.get(old, 0) + n

    # Also walk tables (the address block at top is in a table)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in SIMPLE:
                        n = replace_in_runs(para, old, new)
                        counts[old] = counts.get(old, 0) + n

    # --- Step 5: insert ECI award row ---------------------------------------
    # Find the awards paragraph (contains "Heterocycles Gordon Conference")
    # and prepend a new line. The PDF rendering uses a tab between award
    # title and date to right-align dates; we'll preserve that format.
    eci_text = (
        "ACS Organic Division Early Career Investigator Award – Winner of the "
        "2026 ECI Award\tAug. 2026\n"
    )
    inserted = False
    for para in d.paragraphs:
        text_joined = "".join(r.text for r in para.runs)
        if "Heterocycles Gordon Conference" in text_joined:
            # Prepend ECI row to the same paragraph (preserves tab-stop layout)
            new_text = eci_text + text_joined
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            inserted = True
            break
    counts["+ECI award row"] = 1 if inserted else 0

    d.save(str(DOCX))

    print("=== edit summary ===")
    for k, v in counts.items():
        print(f"  {k!r:60s}  x{v}")


if __name__ == "__main__":
    main()
